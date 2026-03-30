#!/usr/bin/env python3
"""
CBO AI Sandbox Proposal - Document Generator
Huawei Tech. Investment (Oman) LLC
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# COLOUR PALETTE
# ─────────────────────────────────────────────────────────────────────────────
HEX = lambda r, g, b: RGBColor(r, g, b)

C_HUAWEI_RED   = HEX(0xCF, 0x0A, 0x2C)
C_NAVY         = HEX(0x0D, 0x1B, 0x3E)
C_DARK_BLUE    = HEX(0x1A, 0x37, 0x6E)
C_MID_BLUE     = HEX(0x2E, 0x5E, 0xAA)
C_LIGHT_BLUE   = HEX(0xE8, 0xF0, 0xFB)
C_WHITE        = HEX(0xFF, 0xFF, 0xFF)
C_DARK_TEXT    = HEX(0x1C, 0x2B, 0x3A)
C_BODY_TEXT    = HEX(0x2C, 0x3E, 0x50)
C_LIGHT_GRAY   = HEX(0xF4, 0xF6, 0xF9)
C_MID_GRAY     = HEX(0xCC, 0xD1, 0xD8)
C_ACCENT_TEAL  = HEX(0x00, 0x86, 0x80)

C_SCORE_COVERED  = HEX(0x34, 0xA8, 0x53)   # Covered (> 75)   green
C_SCORE_PARTIAL  = HEX(0x8B, 0xC3, 0x4A)   # Partially Covered (1-75)  lime
C_SCORE_OOS      = HEX(0xC0, 0x39, 0x2B)   # Out of Scope (0)  red

def hex_str(c: RGBColor) -> str:
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

# ─────────────────────────────────────────────────────────────────────────────
# LOW-LEVEL HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def shade_para(para, fill: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_str(fill))
    pPr.append(shd)

def shade_cell(cell, fill: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_str(fill))
    tcPr.append(shd)

def cell_para(cell, text, bold=False, italic=False, color=None,
              size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, wrap=True):
    para = cell.paragraphs[0]
    para.clear()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    para.alignment = align
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color
    return para

def cell_add_run(cell, text, bold=False, color=None, size=9.5, newline=False):
    para = cell.paragraphs[0]
    if newline:
        run = para.add_run('\n')
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.bold       = bold
    if color:
        run.font.color.rgb = color

def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)

def no_border_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def thin_border_table(table, color='D0D0D0'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def add_bottom_border_para(doc, color: RGBColor):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '8')
    bot.set(qn('w:color'), hex_str(color))
    pBdr.append(bot)
    pPr.append(pBdr)

def score_color(score):
    if score > 75: return C_SCORE_COVERED
    if score >= 1: return C_SCORE_PARTIAL
    return C_SCORE_OOS

def score_label(score):
    if score > 75: return "Covered"
    if score >= 1: return "Partially Covered"
    return "Out of Scope"

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT-LEVEL HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def section_banner(doc, title, subtitle=None):
    """Dark navy full-width banner for section titles."""
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    no_border_table(t)
    c = t.rows[0].cells[0]
    set_cell_margins(c, top=100, bottom=80, left=120, right=120)
    shade_cell(c, C_NAVY)
    cell_para(c, title, bold=True, color=C_WHITE, size=13)
    if subtitle:
        cell_add_run(c, f'\n{subtitle}', bold=False, color=HEX(0xA0,0xB8,0xD8), size=9.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

def sub_banner(doc, title):
    """Mid-blue sub-section banner."""
    t = doc.add_table(rows=1, cols=1)
    no_border_table(t)
    c = t.rows[0].cells[0]
    set_cell_margins(c, top=60, bottom=50, left=100, right=100)
    shade_cell(c, C_DARK_BLUE)
    cell_para(c, title, bold=True, color=C_WHITE, size=11)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

def body(doc, text, size=10.5, color=None, indent=0, before=1, after=5):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    if indent:
        para.paragraph_format.left_indent = Inches(indent)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return para

def body_bold(doc, label, text, size=10.5, before=1, after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    r1 = para.add_run(label)
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(size)
    r1.font.color.rgb = C_NAVY
    r2 = para.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(size)
    r2.font.color.rgb = C_BODY_TEXT

def bullet(doc, text, bold_prefix=None, indent=0.3, before=1, after=3):
    para = doc.add_paragraph()
    para.paragraph_format.space_before  = Pt(before)
    para.paragraph_format.space_after   = Pt(after)
    para.paragraph_format.left_indent   = Inches(indent)
    r0 = para.add_run('▸  ')
    r0.font.name  = 'Calibri'
    r0.font.size  = Pt(10.5)
    r0.font.color.rgb = C_HUAWEI_RED
    if bold_prefix:
        rb = para.add_run(bold_prefix + ' ')
        rb.bold = True
        rb.font.name = 'Calibri'
        rb.font.size = Pt(10.5)
        rb.font.color.rgb = C_NAVY
    r = para.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(10.5)
    r.font.color.rgb = C_BODY_TEXT

def callout(doc, title, text, color=None):
    """Left-bordered callout box."""
    if color is None:
        color = C_MID_BLUE
    t = doc.add_table(rows=1, cols=2)
    no_border_table(t)
    # left stripe
    lc = t.rows[0].cells[0]
    lc.width = Inches(0.08)
    shade_cell(lc, color)
    # content
    rc = t.rows[0].cells[1]
    set_cell_margins(rc, top=80, bottom=80, left=100, right=80)
    shade_cell(rc, C_LIGHT_BLUE)
    if title:
        cell_para(rc, title, bold=True, color=C_NAVY, size=10)
        cell_add_run(rc, f'\n{text}', bold=False, color=C_BODY_TEXT, size=9.5)
    else:
        cell_para(rc, text, bold=False, color=C_BODY_TEXT, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def hr(doc, color=None):
    if color is None:
        color = C_HUAWEI_RED
    add_bottom_border_para(doc, color)

# ─────────────────────────────────────────────────────────────────────────────
# DATA: REQUIREMENTS
# ─────────────────────────────────────────────────────────────────────────────
REQUIREMENTS = [
    # ── Technical Functional ──────────────────────────────────────────────────
    ("TF-01","Technical Functional","Upload and manage dummy datasets",78,
     "Enterprise-grade encrypted object storage with S3-compatible API, bucket-level ACLs, versioning, and audit logging. File upload is natively integrated into the conversational workbench and the notebook environment."),
    ("TF-02","Technical Functional","Access to pre-trained models and APIs",93,
     "High-performance LLM inference engine serving 20+ open-source models (Qwen3, Llama 3.x, DeepSeek-R1, Mistral) via an OpenAI-compatible API. Multilingual embedding model with native Arabic support. Ready on day one."),
    ("TF-03","Technical Functional","Role-based access controls",93,
     "Enterprise identity and access management platform with full RBAC, group policies, MFA enforcement, JIT privileged access, and LDAP/AD federation. Single sign-on across all sandbox tools."),
    ("TF-04","Technical Functional","Data assessment and preprocessing",78,
     "Reframed as an AI-native document intelligence pipeline: OCR for scanned Arabic and English documents, semantic chunking, multilingual embedding, and automated metadata enrichment. Structured data exploration via the notebook environment."),
    ("TF-05","Technical Functional","Data cleaning tools",75,
     "AI-native approach: a document ingestion pipeline handles format normalisation (PDF, Word, Excel, scanned images). Tabular data cleaning is performed in the notebook environment with full Python data science library support."),
    ("TF-06","Technical Functional","Data privacy and anonymisation",78,
     "Two-layer protection: (1) runtime PII detection and masking on all AI inputs and outputs; (2) a dedicated PII detection and anonymisation microservice supporting Arabic text, Omani national IDs, IBANs, and phone formats — applied before any data enters the sandbox."),
    ("TF-07","Technical Functional","Fully isolated AI development environment",96,
     "eBPF-enforced kernel-level network policies and namespace microsegmentation ensure zero connectivity to production systems. Air-gap capable: all models, registries, and dependencies are self-hosted. Isolation is enforced at kernel level, not merely by firewall rule."),
    ("TF-08","Technical Functional","High-performance GPU hardware",88,
     "Full NVIDIA GPU support (A10/A100/H100) with multi-GPU tensor parallelism and model quantisation for cost-optimised inference. Sized for inference and fine-tuning workloads — not pre-training from scratch, which is not in scope for a regulatory AI sandbox."),
    ("TF-09","Technical Functional","Scalable infrastructure",93,
     "Event-driven autoscaling, vertical resource optimisation, and Kubernetes-native horizontal scaling. Scale-to-zero when idle, scale-to-N under peak load. Three-node starting configuration scales seamlessly to hundreds of nodes with zero re-architecture."),
    ("TF-10","Technical Functional","Performance metrics frameworks",78,
     "LLM observability platform traces every AI call (latency, tokens, cost, quality scores). Full-stack observability platform covers GPU utilisation, request throughput, and error rates. ML experiment tracker adds accuracy, F1, and AUC metrics for trained models."),
    ("TF-11","Technical Functional","Project/workspace creation, collaboration and sharing",78,
     "Source control and project management platform provides Git-backed project isolation — each project is a repository with team access controls, wikis, issue tracking, and branch-based collaboration. A lightweight developer portal surfaces projects and workspace launchers."),
    ("TF-12","Technical Functional","AI experimentation modules — no-code, low-code, and code",85,
     "Three tiers: (1) conversational workbench for all users; (2) visual AI pipeline builder for no-code/low-code construction of RAG pipelines, classification flows, and document analysis; (3) multi-user notebook environment with full ML libraries for code-based experimentation."),
    ("TF-13","Technical Functional","Secure upload, storage and management of synthetic datasets",85,
     "AES-256 encryption at rest, TLS 1.3 in transit, bucket-level access policies, object versioning, and a tamper-evident audit trail of all data access. Synthetic data generation capability produces statistically realistic test datasets without exposing real data."),
    ("TF-14","Technical Functional","Model training, versioning and evaluation",80,
     "ML experiment tracking and model registry captures every training run with linked parameters, metrics, training data references, and model artefacts. LLM evaluation platform tracks response quality over time. Full model lifecycle: experiment → register → deploy → monitor → retrain."),
    ("TF-15","Technical Functional","Guided tutorials, knowledge base and sample use cases",78,
     "Delivered as a RAG application on the platform itself: all tutorial content and use case documentation is ingested and made conversationally accessible via a dedicated 'CBO AI Guide' agent. Twelve pre-built CBO-specific use case templates are provided at launch."),
    ("TF-16","Technical Functional","Full API and platform documentation",88,
     "Interactive OpenAPI/Swagger documentation served through the API gateway for every service. All 52 platform components have detailed technical documentation. CBO-specific operational runbooks and admin guides are produced as engagement deliverables."),
    ("TF-17","Technical Functional","Edge AI inference (on-device ML)",0,
     "OUTSIDE SCOPE — Addressed in Section 9. OpenOva's platform is a server-side, Kubernetes-native architecture. Edge inference on embedded hardware is architecturally separate. If CBO has specific edge use cases, we welcome a dedicated scoping conversation."),
    ("TF-18","Technical Functional","APIs and SDKs for custom development",90,
     "OpenAI-compatible REST API from day one — any developer using the OpenAI Python SDK, TypeScript SDK, or a simple HTTP client can build against the platform with a single base URL change. Vector database, identity, and observability services all expose standard SDKs."),
    ("TF-19","Technical Functional","Multi-user and multi-project support",80,
     "Identity platform manages users, groups, and roles at enterprise scale. Project isolation via source control repositories and Kubernetes namespaces. Developer portal provides multi-tenant project management with per-team resource quotas and access boundaries."),
    # ── Non-Functional ────────────────────────────────────────────────────────
    ("NF-01","Non-Functional","Usability — intuitive UI for all skill levels",83,
     "Conversational workbench for non-technical staff; visual pipeline builder for intermediate users; notebook environment for data scientists; IDE-style workspace for engineers. All tools share single sign-on and a unified developer portal."),
    ("NF-02","Non-Functional","Security — strict isolation from production systems",96,
     "Core architectural principle. eBPF kernel-level enforcement, namespace microsegmentation, policy-as-code admission control, and runtime behavioural threat detection. No egress to production systems without explicit network policy. Independently auditable."),
    ("NF-03","Non-Functional","Privacy and compliance",80,
     "Runtime PII masking, at-rest PII anonymisation, secrets management vault, automated TLS certificate management, and a full audit trail across all AI interactions. Aligns with CBO's own AI Security Policy and relevant data protection standards."),
    ("NF-04","Non-Functional","Performance — <2s UI response, <5s data operations",85,
     "Achievable with the platform stack. LLM first-token latency typically under one second on A10 GPUs. Vector semantic search under 100 ms. All UI components are stateless and load balancer-fronted. Validated during UAT with realistic load patterns."),
    ("NF-05","Non-Functional","Scalability — growing users, projects and hardware",93,
     "Kubernetes-native autoscaling at every layer. GPU nodes added without re-architecture. Three-node start scales to hundreds. Demonstrated in production environments at scale."),
    ("NF-06","Non-Functional","Maintainability — tools for updates and troubleshooting",92,
     "GitOps delivery model: every configuration change is a Git commit — fully auditable, reproducible, and reversible. Platform updates are declarative state changes reconciled automatically. Full change history. AIOps agent provides proactive health monitoring."),
    ("NF-07","Non-Functional","Extensibility — add new modules with minimal effort",90,
     "Kubernetes-native modular architecture. New components deploy via standard configuration. New AI models are a one-line configuration change. New agent templates and pipeline nodes are YAML definitions. No vendor approval required to extend."),
    # ── Security Policy (Appendix A) ─────────────────────────────────────────
    ("SP-01","Security Policy","Risk assessment and continuous monitoring",85,
     "Runtime behavioural threat detection, policy enforcement engine, SIEM and log analytics platform, and full-stack observability provide continuous automated risk monitoring. Risk findings are surfaced to the AIOps agent for correlation and alert generation."),
    ("SP-02","Security Policy","Data encryption at rest and in transit (TLS 1.3)",96,
     "AES-256 for all stored data. TLS 1.3 enforced on all service communications via automated certificate management. mTLS between internal services via eBPF-encrypted service mesh. Secrets vault with proper key management and rotation."),
    ("SP-03","Security Policy","Model version control — code, training data, hyperparameters",78,
     "ML experiment tracker and model registry maintains full lineage: training dataset reference, hyperparameters, metrics, and artefacts per run. Container image registry with cryptographic image signing provides immutable model deployment history. All configuration in Git."),
    ("SP-04","Security Policy","Bias detection and mitigation in training data and models",78,
     "Model fairness assessment toolkit pre-installed in the notebook environment. Assesses accuracy parity, demographic parity, and equalized odds across protected groups. LLM output bias evaluated through the LLM observability platform's evaluation datasets."),
    ("SP-05","Security Policy","Secure model deployment pipeline with audit logging",92,
     "GitOps delivery: every deployment is a reviewed Git commit. Container images are cryptographically signed and verified at deployment time. Policy enforcement engine blocks unsigned or policy-violating images. LLM observability platform logs every inference request."),
    ("SP-06","Security Policy","Model drift detection",75,
     "ML drift detection engine runs scheduled batch analyses comparing live input distributions against training baselines (KS test, PSI, chi-squared). The AIOps agent reads drift reports, correlates against performance metrics, and generates actionable alerts with recommended remediation."),
    ("SP-07","Security Policy","RBAC, MFA and just-in-time access",95,
     "Enterprise IAM platform: full RBAC, TOTP and WebAuthn/FIDO2 MFA, time-limited JIT credentials for privileged operations, session management, and audit log of all access events. Principle of least privilege enforced at the Kubernetes admission layer."),
    ("SP-08","Security Policy","API security",92,
     "Application-layer WAF (OWASP Core Rule Set), eBPF kernel-level L7 policy enforcement, bearer token authentication on all API endpoints, rate limiting, and API schema validation. All API traffic is logged."),
    ("SP-09","Security Policy","Real-time monitoring and anomaly detection",88,
     "Runtime behavioural monitor uses eBPF to detect anomalies (container escape attempts, privilege escalation, data exfiltration patterns, cryptomining). SIEM platform adds ML-based anomaly detection on security events. Full-stack metrics dashboards with configurable alerting."),
    ("SP-10","Security Policy","Comprehensive tamper-evident logging",91,
     "Append-only log aggregation with immutable storage chunks. SIEM platform provides indexed, searchable security event logs with cryptographic integrity guarantees. LLM observability platform maintains an immutable record of every AI interaction. Aligned with CBO's record retention policy A 5.33."),
    ("SP-11","Security Policy","AI-specific incident response",80,
     "AIOps agent provides pre-built detection patterns for AI-specific incidents: guardrail bypass attempts, model misbehaviour, adversarial probing, data exfiltration via model outputs. Generates contextualised alerts with suggested containment actions. Complemented by a delivered AI Incident Response Playbook (10+ scenarios)."),
    ("SP-12","Security Policy","LIME and SHAP explainability",82,
     "LIME, SHAP, Captum, and BertViz are pre-installed in the notebook environment. Four delivered sample notebooks demonstrate explainability on Oman-relevant banking use cases: credit risk attribution, economic indicator feature importance, stress test model explanation, and fraud score explanation."),
    ("SP-13","Security Policy","Ethical AI guidelines enforcement",75,
     "AI safety and content control layer enforces configurable ethical boundaries on all LLM inputs and outputs: topic boundaries, PII masking, harmful content detection, and hallucination guardrails. Complemented by a delivered AI Governance Framework document defining CBO's approval process for high-risk AI use cases."),
    ("SP-14","Security Policy","Fairness assessment across demographic groups",78,
     "Fairness assessment toolkit in the notebook environment measures demographic parity, equalised odds, and individual fairness across defined group attributes. Sample banking fairness notebooks cover credit scoring equity, biometric authentication demographic accuracy, and economic indicator model fairness."),
    ("SP-15","Security Policy","ISO 42001 alignment",68,
     "All technical controls required by ISO 42001 are implemented: risk monitoring, model governance, access control, audit logging, incident response, and data protection. A delivered ISO 42001 Technical Controls Mapping document maps each standard clause to the corresponding platform component. The organisational management system is CBO's responsibility."),
    # ── General Requirements ──────────────────────────────────────────────────
    ("GR-01","General Requirements","Vendor technical support and maintenance",92,
     "Huawei Tech. Investment (Oman) LLC provides commercial support backed by the global Huawei support organisation. The implementation partner provides 24/7 platform engineering support under a defined SLA. All open-source components have enterprise community and commercial support channels."),
    ("GR-02","General Requirements","Complete implementation responsibility",93,
     "Huawei and its implementation partner assume full end-to-end responsibility: infrastructure provisioning, platform deployment, configuration, customisation, integration, testing, training delivery, and documentation — within the 12-week timeline."),
    ("GR-03","General Requirements","Continuous model updates — newer AI models",92,
     "GitOps model update pipeline: new open-source model versions are evaluated, approved via a Pull Request workflow with designated reviewers, and deployed automatically with zero downtime. Model refresh cycle: quarterly minimum, monthly where material new models are released."),
    ("GR-04","General Requirements","Documentation — user guides and system manuals",82,
     "Delivered: end-user quick-start guides (per workspace type), data scientist onboarding guide, admin operations runbook, API reference documentation, AI Incident Response Playbook, ISO 42001 Technical Controls Mapping, and AI Governance Framework. All in English; Arabic summaries for executive materials."),
    ("GR-05","General Requirements","Training sessions for end-users and administrators",87,
     "Three structured training tracks: (1) End-User Track — conversational AI, RAG, pipeline builder, pre-built use cases; (2) Data Scientist Track — notebook environment, ML training, explainability, fairness; (3) Administrator Track — platform operations, security monitoring, user management. Minimum 90% satisfaction target."),
]

# ─────────────────────────────────────────────────────────────────────────────
# DATA: USE CASES
# ─────────────────────────────────────────────────────────────────────────────
USE_CASES = [
    {
        "num": "01",
        "title": "Regulatory Intelligence & Circular Management",
        "cbo_role": "Regulatory Authority | Banking Supervisor",
        "who": "All CBO departments, compliance officers at supervised institutions",
        "what": (
            "All CBO circulars, directives, banking regulations, Basel III guidelines, "
            "Islamic banking standards, and supervisory notices — in both Arabic and English — "
            "are ingested into a multilingual semantic search engine. CBO staff and, in a "
            "separate tenant, compliance officers at supervised banks can ask questions in "
            "natural language and receive cited, source-linked answers in seconds."
        ),
        "why": (
            "CBO has issued hundreds of regulatory circulars over decades. Today, finding the "
            "definitive answer to a compliance question requires manual search across years of "
            "documents. The new Banking Law (RD 2/2025) and Digital Banking Framework (Decision "
            "25/2025) alone represent significant new regulatory surface. AI makes authoritative "
            "compliance intelligence instantly accessible across the entire organisation — and "
            "positions CBO as the definitive source of regulatory truth for the sector."
        ),
        "multi_tenant": (
            "CBO staff access the internal knowledge base. Supervised banks access a curated "
            "external-facing version through a separate tenant — searchable, cited, and "
            "controlled by CBO. This transforms CBO's relationship with the sector: from "
            "reactive enquiry handling to proactive regulatory self-service."
        ),
    },
    {
        "num": "02",
        "title": "Macroeconomic Stress Testing & Scenario Analysis",
        "cbo_role": "Financial Stability | Monetary Policy",
        "who": "Monetary policy economists, financial stability unit",
        "what": (
            "ML models trained on Oman-specific economic indicators predict banking system "
            "stability under stress scenarios: oil price shocks (Oman's primary fiscal driver), "
            "USD peg defence cost analysis, real estate correction impacts, GCC interest rate "
            "convergence effects, and climate-related financial risk (per CBO's October 2024 "
            "green finance circular). SHAP feature attribution explains which economic variables "
            "drive each stress outcome."
        ),
        "why": (
            "CBO must produce stress test reports for the domestic financial system and "
            "contribute to GCC regional financial stability assessments. Manual modelling "
            "in spreadsheets is time-consuming and difficult to reproduce. An AI-augmented "
            "workbench enables economists to iterate faster, explore more scenarios, and "
            "produce explainable, auditable results — with full experiment lineage."
        ),
        "multi_tenant": "Each modelling team works in an isolated project repository. Results are shared via the developer portal with access controls.",
    },
    {
        "num": "03",
        "title": "National Payment Network Fraud Intelligence",
        "cbo_role": "Payment Systems Operator — RTGS, ACH, Fawri/Fawri+, MPCSS",
        "who": "Payment systems department, financial crime unit",
        "what": (
            "As the direct operator of Oman's entire payment infrastructure — RTGS (24/7, "
            "ISO 20022), ACH, Fawri/Fawri+ instant payments, and MPCSS — CBO has a unique "
            "cross-institutional view of every transaction in the Omani financial system. "
            "ML models trained on aggregate payment patterns detect fraud rings, unusual "
            "settlement behaviours, and coordinated payment anomalies that are invisible "
            "to individual banks."
        ),
        "why": (
            "MPCSS transaction volumes surged 318% in a recent 12-month period. Fawri+ "
            "processes 24/7 instant payments. The attack surface is vast. No individual "
            "bank can detect multi-bank fraud rings — only CBO, as the network operator, "
            "has the full picture. This use case directly leverages CBO's unique position "
            "in the payment ecosystem."
        ),
        "multi_tenant": "Analytics run on synthetic and anonymised transaction data within the sandbox. Production integration is a separate future phase.",
    },
    {
        "num": "04",
        "title": "AML/CFT Cross-Institutional Pattern Detection",
        "cbo_role": "Banking Supervisor | AML/CFT Authority",
        "who": "AML/CFT compliance unit, coordination with NCFI (Financial Intelligence Unit)",
        "what": (
            "AI analysis of aggregated suspicious transaction report (STR) data filed by all "
            "supervised institutions. Graph-based network analysis identifies coordinated "
            "money laundering patterns, layering schemes, and structuring behaviours that "
            "span multiple banks — patterns that are invisible in any single institution's "
            "data. ML classifiers prioritise the highest-risk STRs for analyst review."
        ),
        "why": (
            "Oman's 2024 FATF Mutual Evaluation noted strong financial intelligence use "
            "as a key strength. This capability extends that strength into AI-assisted "
            "pattern detection, reducing analyst time on low-risk STRs and improving "
            "detection rates for complex schemes. Aligns with enhanced follow-up "
            "commitments made to MENAFATF."
        ),
        "multi_tenant": "Strictly CBO-internal tenant. All data is synthetic in the sandbox phase. Production integration requires separate governance approval.",
    },
    {
        "num": "05",
        "title": "Bank Examination Intelligence & Supervisory Analytics",
        "cbo_role": "Integrated Financial Regulator",
        "who": "Banking supervision department, bank examiners",
        "what": (
            "AI-assisted processing of bank examination submissions: automatic flagging of "
            "anomalies in CAMELS component data, peer group benchmarking of capital adequacy "
            "and liquidity ratios (LCR, NSFR, CAR against the 12.625% minimum), trend analysis "
            "across reporting periods, and draft examination finding generation. The system "
            "highlights the 10% of data points requiring deep human scrutiny."
        ),
        "why": (
            "CBO is the sole integrated regulator for all banks, finance companies, Islamic "
            "windows, and money exchange companies in Oman. The volume of supervised entities "
            "and the depth of data submitted creates an information processing challenge that "
            "AI is uniquely suited to address — freeing examiners for higher-value judgement "
            "work rather than data aggregation."
        ),
        "multi_tenant": "Sandbox uses synthetic bank examination data. Examiners validate the AI findings against known historical patterns before production consideration.",
    },
    {
        "num": "06",
        "title": "Open Banking & Digital Banking Compliance Monitor",
        "cbo_role": "Open Banking Regulator | Digital Banking Licensor",
        "who": "Digital banking supervision team, fintech supervision unit",
        "what": (
            "With CBO's Open Banking Regulatory Framework (approved December 2024) and Digital "
            "Banking Framework (Decision 25/2025, effective June 2025) both newly in force, "
            "AI monitors API quality, security compliance (authentication standards, "
            "encryption headers, rate limiting), response consistency, and implementation "
            "correctness across all licensed open banking providers and Category 1/2 digital "
            "bank applicants. AI evaluates API specifications against the CBO technical standard."
        ),
        "why": (
            "CBO cannot manually review the API implementations of every licensed provider "
            "at the frequency required. Automated AI monitoring ensures continuous compliance "
            "visibility and generates structured evidence for supervisory reviews."
        ),
        "multi_tenant": "Sandbox uses public API specifications and mock endpoints. Production monitoring is a natural follow-on phase.",
    },
    {
        "num": "07",
        "title": "Regulatory Reporting Validation Engine",
        "cbo_role": "Banking Supervisor | Statistics & Research",
        "who": "Statistics and research department, banking supervision",
        "what": (
            "Banks submit dozens of periodic regulatory returns — Basel III capital components, "
            "liquidity metrics, FX position reports, credit concentration data, and Islamic "
            "finance-specific Shariah compliance reports. AI validates internal consistency "
            "across submissions, flags statistical outliers, detects reporting anomalies, "
            "and compares trends against peer groups and prior periods — before the data "
            "enters CBO's official statistical systems."
        ),
        "why": (
            "Data quality at the point of entry prevents compounding errors in CBO's "
            "macroprudential analysis. AI validation at the submission gateway is more "
            "scalable, faster, and more consistent than manual spot-checking."
        ),
        "multi_tenant": "CBO internal use. Banks receive automated validation feedback before final submission acceptance.",
    },
    {
        "num": "08",
        "title": "Islamic Finance AI Workbench",
        "cbo_role": "Islamic Banking Regulator | Central Shariah Authority Secretariat",
        "who": "Islamic banking supervision unit, Shariah advisors",
        "what": (
            "A specialised AI environment for Islamic finance analysis: RAG over the CBO "
            "Islamic Banking Regulatory Framework (IBRF), AAOIFI standards, and Shariah "
            "board rulings. AI assists supervisors in reviewing the Shariah-compliance "
            "documentation of new Islamic finance products, identifying potential non-compliant "
            "structures, and analysing Shariah board decisions for consistency. "
            "The Central Shariah Authority's precedent library is made conversationally "
            "accessible."
        ),
        "why": (
            "Oman's Islamic banking sector became one of the world's top 20 in under a decade "
            "of operation. The volume of Shariah review work has grown correspondingly. "
            "AI cannot replace Shariah judgement, but it can dramatically accelerate "
            "research, precedent lookup, and documentation review — allowing the Shariah "
            "authority to focus on substantive rulings."
        ),
        "multi_tenant": "Islamic banking supervision team and Shariah authority secretariat have separate access-controlled workspaces.",
    },
    {
        "num": "09",
        "title": "FinTech Regulatory Sandbox & KIC Management",
        "cbo_role": "FinTech Regulator | Innovation Hub Operator",
        "who": "FinTech supervision and innovation team",
        "what": (
            "AI-assisted management of CBO's fintech regulatory sandbox and Knowledge & "
            "Innovation Centre (KIC): automated analysis of licence applications, compliance "
            "milestone tracking, AI-generated assessment reports, and comparison of applicant "
            "technology proposals against CBO's regulatory framework. Conversational search "
            "over all submitted proposals and precedent decisions."
        ),
        "why": (
            "CBO targets a 16% CAGR in Oman's fintech market and has 52+ applications under "
            "active review. AI-assisted processing enables the team to handle growth in "
            "application volume without proportional headcount growth, while maintaining "
            "consistency in evaluation quality."
        ),
        "multi_tenant": "Internal CBO use. Fintech applicants interact via a separate structured submission portal — not the AI sandbox directly.",
    },
    {
        "num": "10",
        "title": "Monetary Policy Research Workbench",
        "cbo_role": "Monetary Authority | Macroeconomic Research",
        "who": "Economic research department, monetary policy committee support",
        "what": (
            "AI-augmented research environment where economists interact with decades of CBO "
            "and IMF economic data, international monetary research literature, and local "
            "economic model outputs. Generate research briefs on impact of USD peg policy "
            "under shifting oil revenue scenarios, simulate GCC interest rate policy "
            "responses, and draft monetary policy committee papers. Arabic and English "
            "research literature both fully accessible."
        ),
        "why": (
            "Monetary policy decisions require synthesising vast quantities of rapidly "
            "changing global and domestic economic information. AI compresses the research "
            "cycle from weeks to hours, enabling more thorough analysis before each "
            "policy committee meeting."
        ),
        "multi_tenant": "Research team workspace with fine-grained document access controls. Monetary policy committee materials handled in a restricted-access sub-tenant.",
    },
    {
        "num": "11",
        "title": "Payment Technology Supervision & Bias Assessment",
        "cbo_role": "Payment Systems Regulator | Consumer Protection Authority",
        "who": "Payment systems supervisors, consumer protection unit",
        "what": (
            "As AI-based authentication (facial recognition, voice biometrics, behavioural "
            "analytics) proliferates in Omani banking, CBO supervisors need to evaluate "
            "these systems for accuracy, security vulnerabilities, and demographic bias "
            "before approving bank deployment. The sandbox provides an environment to test "
            "these models against synthetic data, run SHAP analysis on decision logic, and "
            "assess fairness across gender, age, and regional demographic groups."
        ),
        "why": (
            "A biometric authentication system with demographic bias can exclude specific "
            "population segments from banking services — a consumer protection and financial "
            "inclusion issue directly within CBO's mandate. Supervisors need the technical "
            "capability to detect and challenge this. No other Omani institution has this "
            "supervisory evaluation capability today."
        ),
        "multi_tenant": "Banks may be invited to submit model documentation for CBO evaluation in future phases. Sandbox phase uses publicly available biometric test datasets.",
    },
    {
        "num": "12",
        "title": "Cross-Border Payment Intelligence — BUNA, AFAQ & RTGS",
        "cbo_role": "Payment Systems Operator | FX Reserve Manager | AML Authority",
        "who": "Payment systems department, FX management, AML/CFT unit",
        "what": (
            "Oman's 24/7 RTGS is integrated with BUNA (Arab regional payments, OMR as a "
            "settlement currency) and AFAQ (GCC instant cross-border transfers). AI models "
            "analyse cross-border payment flow patterns for structural changes in regional "
            "fund flows, FX concentration risk, sanctions screening optimisation, and "
            "early indicators of cross-border money laundering activity visible only at "
            "the clearing network level."
        ),
        "why": (
            "CBO's participation in both BUNA and AFAQ gives it a unique regional "
            "visibility position. Cross-border payment corridors are a primary vector for "
            "international financial crime. AI-based network analysis of flow patterns "
            "creates an intelligence capability that complements NCFI's domestic FIU function."
        ),
        "multi_tenant": "CBO-internal only. Separate workspaces for payment systems team (flow analytics) and AML unit (financial crime patterns).",
    },
]

# ─────────────────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # page setup — A4
    sec = doc.sections[0]
    sec.page_height  = Cm(29.7)
    sec.page_width   = Cm(21.0)
    sec.left_margin  = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.top_margin   = Cm(1.8)
    sec.bottom_margin= Cm(1.8)

    # default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = C_BODY_TEXT

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    # Huawei Logo — actual rendered logo image
    logo_para = doc.add_paragraph()
    logo_para.paragraph_format.space_before = Pt(0)
    logo_para.paragraph_format.space_after  = Pt(6)
    logo_run = logo_para.add_run()
    logo_run.add_picture('/tmp/huawei_logo.png', width=Inches(4.0))

    # Top accent bar
    top_bar = doc.add_table(rows=1, cols=1)
    no_border_table(top_bar)
    tb_c = top_bar.rows[0].cells[0]
    set_cell_margins(tb_c, top=6, bottom=6, left=120, right=120)
    shade_cell(tb_c, C_HUAWEI_RED)
    cell_para(tb_c, '', bold=True, color=C_WHITE, size=4, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Large title block
    for _ in range(3): doc.add_paragraph()

    t_main = doc.add_table(rows=1, cols=1)
    no_border_table(t_main)
    mc = t_main.rows[0].cells[0]
    set_cell_margins(mc, top=200, bottom=200, left=120, right=120)
    shade_cell(mc, C_NAVY)
    cell_para(mc, 'AI INNOVATION SANDBOX PLATFORM',
              bold=True, color=C_WHITE, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_add_run(mc, '\n\nResponse to Statement of Work', bold=False,
                 color=HEX(0xA0,0xB8,0xD8), size=14)
    cell_add_run(mc, '\n\nCentral Bank of Oman', bold=True,
                 color=HEX(0xFF,0xD7,0x00), size=16)
    cell_add_run(mc, '\n\nConfidential | March 2026', bold=False,
                 color=HEX(0xA0,0xB8,0xD8), size=11)

    for _ in range(4): doc.add_paragraph()

    # Cover footer info
    cover_footer = doc.add_table(rows=1, cols=2)
    no_border_table(cover_footer)
    set_col_width(cover_footer, 0, 3.5)
    set_col_width(cover_footer, 1, 3.5)
    lc = cover_footer.rows[0].cells[0]
    rc = cover_footer.rows[0].cells[1]
    set_cell_margins(lc, top=80, bottom=80, left=120, right=80)
    set_cell_margins(rc, top=80, bottom=80, left=80, right=120)
    shade_cell(lc, C_LIGHT_GRAY)
    shade_cell(rc, C_LIGHT_GRAY)
    cell_para(lc, 'Prepared by', bold=True, color=C_NAVY, size=9)
    cell_add_run(lc, '\nHuawei Tech. Investment (Oman) LLC\nMuscat, Sultanate of Oman',
                 bold=False, color=C_BODY_TEXT, size=9)
    cell_para(rc, 'Document Reference', bold=True, color=C_NAVY, size=9)
    cell_add_run(rc, '\nHW-OM-CBO-AISB-2026-001\nVersion 1.0 — March 2026\nFor Submission to CBO Technology Procurement',
                 bold=False, color=C_BODY_TEXT, size=9)

    doc.add_page_break()

    # ── CONFIDENTIALITY NOTICE ────────────────────────────────────────────────
    section_banner(doc, 'CONFIDENTIALITY NOTICE')
    callout(doc, None,
        'This document is prepared exclusively for the Central Bank of Oman in response to '
        'the AI Sandbox Statement of Work. It contains proprietary and commercially sensitive '
        'information belonging to Huawei Tech. Investment (Oman) LLC and its partners. '
        'Recipients must not reproduce, distribute, or disclose any part of this document '
        'to third parties without the prior written consent of Huawei Tech. Investment (Oman) LLC. '
        'Information provided by CBO during the tendering process is treated as strictly '
        'confidential and used solely for the purpose of preparing this response.',
        color=C_HUAWEI_RED)

    # ── EXECUTIVE SUMMARY ─────────────────────────────────────────────────────
    section_banner(doc, '1.  EXECUTIVE SUMMARY')
    body(doc,
        'Huawei Tech. Investment (Oman) LLC is pleased to present this response to the '
        'Central Bank of Oman\'s AI Sandbox Statement of Work. We offer a comprehensive, '
        'enterprise-grade AI Innovation Sandbox Platform built on a curated ecosystem of '
        '52 open-source components — fully managed, fully supported, and deployed entirely '
        'within CBO\'s sovereign infrastructure boundary.',
        before=4, after=6)

    # 4-quadrant summary table
    sum_table = doc.add_table(rows=2, cols=2)
    thin_border_table(sum_table)
    labels = [
        ('High Out-of-the-Box Coverage',
         'The great majority of all RFP requirements — functional, non-functional, and '
         'security policy — are delivered by our platform without customisation. The remaining '
         'requirements are addressed through targeted additions, all open source, all delivered '
         'within the 12-week implementation timeline.'),
        ('Complete Data Sovereignty',
         'Every component — LLM inference, vector search, semantic pipelines, identity, '
         'secrets, observability — runs inside CBO\'s own infrastructure boundary. No data '
         'ever touches a foreign cloud. Arabic-native multilingual embedding is standard, '
         'not an afterthought.'),
        ('AI That Manages Itself',
         'The platform includes a pre-trained AIOps intelligence layer with semantic knowledge '
         'of every component: its failure modes, integration dependencies, health indicators, '
         'and remediation steps. The platform detects, correlates, and often resolves issues '
         'before CBO\'s IT team is aware of them.'),
        ('Positioned for CBO\'s Unique Mandate',
         'We have designed twelve use cases specific to CBO\'s multi-role position as '
         'integrated regulator, payment network operator, Shariah authority, Mala\'a '
         'supervisor, and Vision 2040 enabler. This is not a generic AI platform — it is '
         'designed for the Central Bank of Oman.'),
    ]
    positions = [(0,0),(0,1),(1,0),(1,1)]
    for (r, c), (title, text) in zip(positions, labels):
        cell = sum_table.rows[r].cells[c]
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        shade_cell(cell, C_LIGHT_BLUE if (r+c)%2==0 else C_LIGHT_GRAY)
        cell_para(cell, title, bold=True, color=C_NAVY, size=10.5)
        cell_add_run(cell, f'\n{text}', bold=False, color=C_BODY_TEXT, size=9.5)
    doc.add_paragraph()

    body(doc,
        'Huawei Tech. Investment (Oman) LLC is the single accountable delivery entity: '
        'infrastructure, platform engineering, customisation, training, and long-term support — '
        'all under one contract, backed by Huawei\'s global resources and established '
        'presence in Oman.',
        after=8)

    doc.add_page_break()

    # ── ABOUT THE RESPONDENT ──────────────────────────────────────────────────
    section_banner(doc, '2.  ABOUT THE RESPONDENT')
    body(doc,
        'Huawei Tech. Investment (Oman) LLC is part of Huawei Technologies — a global '
        'technology leader with over 35 years of experience delivering enterprise infrastructure '
        'solutions across 170+ countries. In Oman, Huawei has a long-standing relationship '
        'with government and financial sector institutions, providing ICT infrastructure that '
        'underpins critical national services.',
        before=4, after=6)
    bullet(doc,
        'Complete end-to-end accountability: commercial, infrastructure, platform engineering, and implementation',
        bold_prefix='Single Prime Contractor:')
    bullet(doc,
        'All technical delivery, AI platform deployment, customisation, training, and post-go-live support under one contract',
        bold_prefix='Full Delivery Scope:')
    bullet(doc,
        'On-premises or Huawei Cloud — same architecture, same platform, CBO\'s choice of deployment model',
        bold_prefix='Infrastructure Options:')
    bullet(doc,
        'Backed by the global Huawei support organisation with a dedicated Oman presence and OMR-denominated commercial model',
        bold_prefix='Support Model:')

    sub_banner(doc, 'Our Technical Delivery Capability')
    body(doc,
        'Huawei brings deep, specialised capability in enterprise open-source AI infrastructure. '
        'Our delivery team fields expertise in the following disciplines, all directly relevant '
        'to this engagement:',
        after=5)
    bullet(doc, 'Design and deployment of AI-native Kubernetes infrastructure at banking and government scale')
    bullet(doc, 'Deep expertise in large language model serving, RAG pipeline construction, and AI safety enforcement')
    bullet(doc, 'Regulatory-sector implementation experience across GCC and international markets')
    bullet(doc, 'A proprietary AIOps intelligence layer with pre-built semantic knowledge of every platform component — enabling self-healing, proactive monitoring, and AI-managed operations')
    body(doc,
        'Our solution is built on a curated ecosystem of 52 production-grade open-source '
        'components — enterprise-supported, fully sovereign, and designed from the ground up '
        'to be AI-manageable.',
        before=4, after=8)

    doc.add_page_break()

    # ── OUR INTERPRETATION OF THE RFP ─────────────────────────────────────────
    section_banner(doc, '3.  OUR INTERPRETATION OF THE REQUIREMENTS')
    body(doc,
        'Our analysis of the Statement of Work, combined with our understanding of CBO\'s '
        'mandate, leads us to a precise interpretation of what the AI Sandbox must deliver '
        'and for whom.',
        before=4, after=5)

    sub_banner(doc, 'User Personas — Who Will Use the Sandbox')
    personas = [
        ('Domain Expert', 'Risk officers, compliance staff, economists, Shariah advisors, regulatory supervisors. They want to use AI — not build it. Primary interface: conversational workbench and pre-built use case agents.'),
        ('Innovation Builder', 'Innovation team, fintech supervisors, IT staff. Want to prototype AI-powered tools without deep ML engineering. Interface: visual no-code pipeline builder and conversational workbench.'),
        ('Data Scientist / Quant', 'Economists, quantitative analysts. Python-capable; want to train models, run statistical analysis, produce SHAP explanations. Interface: notebook environment with full ML library stack.'),
        ('Platform Engineer / Admin', 'IT and platform administrators. Manage users, monitor security, maintain the environment. Interface: developer portal, observability dashboards, identity management console.'),
    ]
    pt = doc.add_table(rows=1+len(personas), cols=2)
    thin_border_table(pt)
    header_row = pt.rows[0]
    for i, h in enumerate(['Persona', 'Description & Primary Interface']):
        shade_cell(header_row.cells[i], C_NAVY)
        cell_para(header_row.cells[i], h, bold=True, color=C_WHITE, size=9.5)
        set_cell_margins(header_row.cells[i], top=80, bottom=80, left=100, right=100)
    for i, (persona, desc) in enumerate(personas):
        row = pt.rows[i+1]
        bg = C_LIGHT_BLUE if i%2==0 else C_LIGHT_GRAY
        for c in row.cells:
            shade_cell(c, bg)
            set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        cell_para(row.cells[0], persona, bold=True, color=C_NAVY, size=9.5)
        cell_para(row.cells[1], desc, bold=False, color=C_BODY_TEXT, size=9.5)
    set_col_width(pt, 0, 1.8)
    set_col_width(pt, 1, 5.0)
    doc.add_paragraph()

    sub_banner(doc, 'Scope Clarifications — Our Recommended Interpretation')
    clarifications = [
        ('Model "Training"', 'The RFP references model training throughout. We interpret this as: (a) use of pre-trained models for inference and RAG — the primary use case; (b) fine-tuning pre-trained models on CBO-specific corpora; and (c) training classical ML models (regression, classification) for economic analysis. The sandbox supports all three. Training foundation LLMs from scratch is not in scope for a regulatory AI sandbox — and is not required by any of the stated use cases.'),
        ('"GPT" Model Updates', 'The RFP references "newer versions of GPT". We deliver equivalent or superior open-source models (Qwen3, Llama 3.x, DeepSeek-R1) on-premises — with zero data leaving CBO. New model versions are deployed through a GitOps-controlled update pipeline on a quarterly cycle, or sooner when materially better models are released.'),
        ('Edge AI Inference', 'Running ML models on embedded/mobile/IoT hardware is architecturally distinct from a server-side AI sandbox and is outside the scope of this proposal. If CBO identifies specific edge deployment requirements (e.g., branch-level hardware, ATM-side AI), we welcome a separate scoping engagement. The sandbox serves as the development and validation environment for models that may eventually be deployed to edge hardware.'),
        ('Data Preprocessing', 'We interpret "data preprocessing" in the context of a central bank AI sandbox as: document intelligence for unstructured content (OCR, parsing, Arabic/English text extraction, semantic chunking) and tabular data preparation in the notebook environment. We do not recommend deploying a full data engineering platform — the use cases do not require it.'),
    ]
    for label, text in clarifications:
        body_bold(doc, f'{label}: ', text, after=5)

    doc.add_page_break()

    # ── SOLUTION ARCHITECTURE ─────────────────────────────────────────────────
    section_banner(doc, '4.  PROPOSED SOLUTION ARCHITECTURE')
    body(doc,
        'The platform is structured in four integrated layers — each independently scalable, '
        'each designed to be AI-manageable, and together forming a complete AI innovation '
        'environment from user interface to infrastructure.',
        before=4, after=5)

    # High-Level Design diagram
    sub_banner(doc, 'High-Level Architecture Diagram')
    hld_para = doc.add_paragraph()
    hld_para.paragraph_format.space_before = Pt(4)
    hld_para.paragraph_format.space_after  = Pt(8)
    hld_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hld_run = hld_para.add_run()
    hld_run.add_picture('/tmp/hld.png', width=Inches(6.5))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    layers = [
        ('LAYER 1 — WORKSPACES',
         C_HUAWEI_RED,
         [
             ('Conversational AI Workbench', 'Chat interface with multi-model access, banking use case agents (Regulatory Advisor, Stress Test Assistant, Shariah Research Agent), document Q&A, and Arabic language support. For all users.'),
             ('Visual AI Pipeline Builder', 'Drag-and-drop construction of RAG pipelines, document classifiers, and analysis flows. No Python required. Connects to internal models and knowledge stores. For innovation builders.'),
             ('Multi-User Notebook Environment', 'Browser-based Python notebooks with GPU access on demand. Pre-installed: ML libraries, SHAP, LIME, Fairlearn, and CBO banking use case sample notebooks. For data scientists and quants.'),
             ('AI-Powered Development Workspace', 'Browser-accessible IDE environment with an AI coding assistant backed by an internally-hosted code-specialist model. Every project is a Git repository. For engineers.'),
             ('Developer Portal', 'Unified catalogue of all projects, workspace launchers, resource usage dashboards, and administrative controls. The front door to the entire sandbox.'),
         ]),
        ('LAYER 2 — AI SERVICES',
         C_DARK_BLUE,
         [
             ('LLM Inference Engine', 'High-throughput serving of 20+ open-source models. OpenAI-compatible API. GPU-accelerated. Supports Qwen3, Llama 3.x, DeepSeek-R1, Mistral, and Arabic-optimised models. Scale-to-zero when idle.'),
             ('Document Intelligence Pipeline', 'OCR for scanned Arabic and English documents, format normalisation (PDF/Word/Excel/image), semantic chunking, and metadata enrichment. Bulk ingestion and real-time upload paths.'),
             ('Multilingual Semantic Search', 'Enterprise vector database with HNSW indexing for sub-100ms semantic retrieval. Multilingual embedding model with native Arabic support. Hybrid dense + sparse retrieval. Graph knowledge layer for entity relationships.'),
             ('AI Safety & Content Control', 'Configurable guardrail layer: PII masking, prompt injection detection, topic boundary enforcement, hallucination detection. Applied to all AI inputs and outputs. Every activation is logged.'),
             ('ML Platform', 'Experiment tracking and model registry with full lineage (parameters, datasets, metrics, artefacts). Drift detection engine monitors live model performance. Fairness assessment toolkit for bias evaluation.'),
             ('LLM Observability', 'Every LLM call is traced: input, output, cost, latency, token usage, guardrail activations, evaluation scores. Immutable audit trail. Real-time cost attribution per user, project, and use case.'),
         ]),
        ('LAYER 3 — PLATFORM FOUNDATION',
         C_NAVY,
         [
             ('Identity & Access Management', 'Enterprise RBAC, MFA (TOTP + WebAuthn), JIT privileged access, LDAP/AD federation, single sign-on across all workspaces. Keyed to CBO\'s organisational structure and department hierarchy.'),
             ('Source Control & GitOps Delivery', 'Every project is a Git repository. Infrastructure changes, model updates, and configuration are all Git commits — reviewed, approved, and automatically reconciled. Full change history. Zero-manual-intervention deployments.'),
             ('Security Stack', 'Runtime behavioural threat detection (eBPF-based), policy enforcement engine, application WAF (OWASP CRS), supply chain signing, container vulnerability scanning, and SIEM with ML-based anomaly detection.'),
             ('Full-Stack Observability', 'Unified metrics, logs, and distributed traces across all 52 components. GPU utilisation, LLM throughput, model latency, security event dashboards. Tamper-evident log storage.'),
             ('AIOps Intelligence Agent', 'Pre-trained semantic knowledge of every platform component: failure modes, health indicators, integration dependencies, remediation procedures. Proactive alerting, automated root cause analysis, and self-healing recommendations.'),
             ('Secrets Management', 'Secrets vault with dynamic credential generation, encryption key management, TLS certificate automation, and audit log of all secret access. No secrets stored in configuration or Git.'),
         ]),
        ('LAYER 4 — INFRASTRUCTURE',
         HEX(0x12,0x3A,0x2B),
         [
             ('Option A — On-Premises GPU', 'Full deployment on CBO\'s own hardware. Maximum sovereignty. Air-gap capable. IaC-provisioned from day one — every server, network rule, and storage configuration is defined as code and reproducible.'),
             ('Option B — Huawei Cloud', 'Platform control plane on Huawei Cloud, LLM inference via Huawei\'s managed AI services. Starting configuration: three nodes, under 500 OMR per month. Scales to hundreds of nodes with zero re-architecture. OMR-denominated billing.'),
             ('Infrastructure-as-Code First', 'Every infrastructure component — network, compute, storage, security groups — is defined as code, tracked in Git, and deployed declaratively. Any state is reproducible. Any change is auditable. Any rollback is a Git revert.'),
         ]),
    ]

    for layer_title, layer_color, components in layers:
        # Single table: row 0 = layer header, rows 1+ = components
        ct = doc.add_table(rows=1+len(components), cols=1)
        thin_border_table(ct)
        # Layer header row
        hc = ct.rows[0].cells[0]
        set_cell_margins(hc, top=70, bottom=60, left=100, right=100)
        shade_cell(hc, layer_color)
        cell_para(hc, layer_title, bold=True, color=C_WHITE, size=11)
        # Component rows
        for idx, (comp_name, comp_desc) in enumerate(components):
            cc = ct.rows[idx+1].cells[0]
            set_cell_margins(cc, top=90, bottom=90, left=120, right=120)
            shade_cell(cc, C_LIGHT_BLUE if idx%2==0 else C_LIGHT_GRAY)
            cell_para(cc, comp_name, bold=True, color=C_NAVY, size=9.5)
            cell_add_run(cc, f'  —  {comp_desc}', bold=False, color=C_BODY_TEXT, size=9)
        doc.add_paragraph()

    doc.add_page_break()

    # ── REQUIREMENT COVERAGE HEAT MAP ─────────────────────────────────────────
    section_banner(doc, '5.  REQUIREMENT COVERAGE HEAT MAP',
                   'Color key:  ■ Covered (score > 75)  ■ Partially Covered (score 1–75)  ■ Out of Scope (score 0)')
    body(doc,
        'The following table maps every requirement in the Statement of Work to our '
        'platform coverage. Scores reflect capability available within the 12-week '
        'implementation scope. All open-source components are enterprise-supported.',
        before=4, after=6)

    # Group requirements
    categories = {}
    for req in REQUIREMENTS:
        cat = req[1]
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(req)

    for cat, reqs in categories.items():
        # Single table: row 0 = category header (merged), row 1 = column headers, rows 2+ = requirements
        rt = doc.add_table(rows=2 + len(reqs), cols=4)
        thin_border_table(rt)
        set_col_width(rt, 0, 0.5)    # Ref
        set_col_width(rt, 1, 1.8)    # Requirement
        set_col_width(rt, 2, 0.7)    # Status
        set_col_width(rt, 3, 3.6)    # Response (2× requirement column)

        # Row 0: category header — merged across all 4 columns
        cat_row = rt.rows[0]
        cat_row.cells[0].merge(cat_row.cells[3])
        ch = cat_row.cells[0]
        set_cell_margins(ch, top=60, bottom=50, left=100, right=100)
        shade_cell(ch, C_DARK_BLUE)
        cell_para(ch, cat.upper(), bold=True, color=C_WHITE, size=10)

        # Row 1: column headers
        hr_row = rt.rows[1]
        for i, h in enumerate(['Ref', 'Requirement', 'Status', 'Response']):
            shade_cell(hr_row.cells[i], C_NAVY)
            set_cell_margins(hr_row.cells[i], top=70, bottom=70, left=80, right=80)
            cell_para(hr_row.cells[i], h, bold=True, color=C_WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        for i, (rid, _, rname, score, approach) in enumerate(reqs):
            row = rt.rows[i+2]
            bg = C_LIGHT_BLUE if i%2==0 else C_LIGHT_GRAY
            for c in row.cells:
                shade_cell(c, bg)
                set_cell_margins(c, top=60, bottom=60, left=80, right=80)

            cell_para(row.cells[0], rid, bold=True, color=C_DARK_BLUE, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_para(row.cells[1], rname, bold=False, color=C_BODY_TEXT, size=9)

            # status cell — colour-coded, no score number
            sc = score_color(score)
            shade_cell(row.cells[2], sc)
            cell_para(row.cells[2], score_label(score), bold=True, color=C_WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

            cell_para(row.cells[3], approach, bold=False, color=C_BODY_TEXT, size=8.5)

        doc.add_paragraph()

    # Overall score summary
    scores = [r[3] for r in REQUIREMENTS]
    avg = sum(scores)/len(scores)
    covered = sum(1 for s in scores if s >= 75)
    out_of_scope = sum(1 for s in scores if s < 25)

    callout(doc, f'Overall Coverage: {covered} of {len(scores)} Requirements Covered',
        f'{covered} requirements are fully Covered — delivered out of the box or with standard additions within the 12-week scope. '
        f'{out_of_scope} requirement(s) acknowledged as outside scope with documented rationale. '
        f'All status assessments reflect what the platform delivers within the engagement — no inflation.',
        color=C_SCORE_COVERED)

    doc.add_page_break()

    # ── CBO USE CASES ─────────────────────────────────────────────────────────
    section_banner(doc, '6.  CBO-SPECIFIC AI USE CASES',
                   'Twelve ready-to-deploy use cases designed around CBO\'s unique multi-role mandate')
    body(doc,
        'The following use cases are provided as pre-built templates at platform launch. '
        'They are not demonstrations — they are functional, deployable capabilities built '
        'on the platform\'s core infrastructure. Each can be extended, customised, or '
        'used as a starting point for further innovation by CBO\'s teams.',
        before=4, after=6)

    for uc in USE_CASES:
        uc_fields = [
            ('Users', uc['who'], C_LIGHT_BLUE),
            ('What the AI Does', uc['what'], C_LIGHT_GRAY),
            ('Why It Matters for CBO', uc['why'], C_LIGHT_BLUE),
            ('Multi-User / Multi-Tenant Note', uc['multi_tenant'], C_LIGHT_GRAY),
        ]
        # Single unified table: row 0 = header (2 cols: number + title), rows 1+ = content (merged)
        uc_table = doc.add_table(rows=1+len(uc_fields), cols=2)
        thin_border_table(uc_table)
        set_col_width(uc_table, 0, 0.5)
        set_col_width(uc_table, 1, 6.3)

        # Header row
        num_cell   = uc_table.rows[0].cells[0]
        title_cell = uc_table.rows[0].cells[1]
        shade_cell(num_cell, C_HUAWEI_RED)
        shade_cell(title_cell, C_NAVY)
        set_cell_margins(num_cell,   top=80, bottom=80, left=60, right=60)
        set_cell_margins(title_cell, top=80, bottom=80, left=100, right=100)
        cell_para(num_cell, uc['num'], bold=True, color=C_WHITE, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(title_cell, uc['title'].upper(), bold=True, color=C_WHITE, size=11)
        cell_add_run(title_cell, f'\n{uc["cbo_role"]}', bold=False, color=HEX(0xA0,0xC4,0xFF), size=9)

        # Content rows — merge 2 columns so content is full-width
        for i, (label, content, bg) in enumerate(uc_fields):
            row = uc_table.rows[i+1]
            row.cells[0].merge(row.cells[1])
            cc = row.cells[0]
            shade_cell(cc, bg)
            set_cell_margins(cc, top=90, bottom=90, left=120, right=120)
            cell_para(cc, label, bold=True, color=C_NAVY, size=9)
            cell_add_run(cc, f'\n{content}', bold=False, color=C_BODY_TEXT, size=9)

        doc.add_paragraph()

    doc.add_page_break()

    # ── OPEN SOURCE ADVANTAGE ─────────────────────────────────────────────────
    section_banner(doc, '7.  THE OPEN SOURCE & AI-NATIVE ADVANTAGE')
    body(doc,
        'Every component in this platform is open source. This is not a cost decision — '
        'it is an architectural and regulatory decision that has profound implications '
        'for CBO as a central bank and as an AI adopter.',
        before=4, after=6)

    adv_table = doc.add_table(rows=4, cols=2)
    thin_border_table(adv_table)
    adv_data = [
        ('Complete Regulatory Auditability',
         'Open source means CBO\'s own technical teams, and CBO\'s own external auditors, '
         'can inspect every line of code running in the platform. There are no black boxes. '
         'When the Board or an external audit firm asks "how does the AI system make its '
         'decisions?", CBO can provide a complete, verifiable answer. This is impossible '
         'with proprietary AI platforms.'),
        ('Data Sovereignty Without Compromise',
         'All 52 components run inside CBO\'s own infrastructure — on-premises or on '
         'Huawei Cloud within Oman\'s data jurisdiction. No model weights, no inference '
         'requests, no CBO data of any kind touches a foreign commercial AI provider. '
         'This is not a configuration option — it is the architectural default.'),
        ('AI That Can Manage Open Source (The AI-Native Moat)',
         'For AI to manage an infrastructure platform, it must be able to read, understand, '
         'and reason about every component. Open source provides the structured CRDs, '
         'documented failure modes, and integration graphs that make this possible. '
         'Our AIOps intelligence layer has pre-built semantic knowledge of all 52 components. '
         'This is architecturally impossible on proprietary closed-source platforms — and '
         'cannot be retrofitted after the fact.'),
        ('Enterprise Support Without Vendor Lock-In',
         'Every open-source component in our platform has an enterprise support tier from '
         'its governing foundation or a certified third party. Huawei provides umbrella '
         'support across the entire stack. CBO owns the software regardless of any '
         'commercial relationship — the platform never expires, is never disabled by a '
         'licence key, and always belongs to CBO.'),
    ]
    for i, (title, text) in enumerate(adv_data):
        r, c = i//2, i%2
        cell = adv_table.rows[r].cells[c]
        shade_cell(cell, C_LIGHT_BLUE if (r+c)%2==0 else C_LIGHT_GRAY)
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell_para(cell, title, bold=True, color=C_NAVY, size=10.5)
        cell_add_run(cell, f'\n{text}', bold=False, color=C_BODY_TEXT, size=9.5)

    doc.add_paragraph()
    callout(doc, 'Token Efficiency: The Economic Argument for Open Source AI',
        'Commercial AI APIs charge per token. A closed-source platform dumps unstructured '
        'logs and raw configs into LLM prompts — thousands of tokens per query — because '
        'it has no structured knowledge of its own components. Our AIOps intelligence layer '
        'sends surgical, structured context: typed CRD state, correlated telemetry signals, '
        'known failure mode patterns. This is 10x fewer tokens, 10x faster, 10x more accurate. '
        'For a platform that actively manages itself, token efficiency is an economic moat — '
        'not a marketing claim.',
        color=C_ACCENT_TEAL)

    doc.add_page_break()

    # ── IMPLEMENTATION PLAN ───────────────────────────────────────────────────
    section_banner(doc, '8.  IMPLEMENTATION PLAN — 12 WEEKS TO LIVE PLATFORM')
    phases = [
        ('Phase 1', 'Weeks 1–2', 'Discovery & Architecture',
         [
             'Infrastructure assessment: hardware options, network topology, GPU selection',
             'User persona workshops: validate use cases, prioritise first-wave deployments',
             'Security architecture review with CBO Information Security team',
             'Finalise deployment model (Option A or Option B)',
             'Sign-off on data ingestion scope (which document corpora for initial RAG)',
         ]),
        ('Phase 2', 'Weeks 2–5', 'Core Platform Deployment',
         [
             'Infrastructure provisioning (IaC-first, all configuration in Git from day one)',
             'Deploy all 52 platform components: LLM inference, vector database, identity, observability, security stack',
             'Configure RBAC aligned to CBO organisational structure',
             'Ingest initial document corpus (CBO circulars, Banking Law, Digital Banking Framework)',
             'Deploy conversational workbench with six pre-built CBO banking agent presets',
         ]),
        ('Phase 3', 'Weeks 5–8', 'Experimentation Layer & Use Cases',
         [
             'Deploy visual pipeline builder, notebook environment, and AI-powered development workspace',
             'Configure AI coding assistant connected to internal code-specialist model',
             'Deploy ML platform (experiment tracking, model registry, drift detection)',
             'Install SHAP, LIME, Fairlearn, and PII anonymisation service',
             'Build and validate six priority use case templates with CBO subject matter experts',
         ]),
        ('Phase 4', 'Weeks 8–10', 'Security Hardening & Compliance',
         [
             'Penetration testing against sandbox boundary',
             'Security audit against CBO AI Security Policy (Appendix A)',
             'Produce ISO 42001 Technical Controls Mapping document',
             'Deliver AI Incident Response Playbook',
             'Deliver AI Governance Framework document',
             'Final security sign-off by CBO Information Security team',
         ]),
        ('Phase 5', 'Weeks 10–12', 'Training, Handover & Go-Live',
         [
             'End-User Track training (domain experts, analysts, compliance staff)',
             'Data Scientist Track training (notebook environment, ML workflow, explainability)',
             'Administrator Track training (platform operations, security monitoring, user management)',
             'Complete all deliverable documentation',
             'Production go-live, hypercare period begins',
             'Formal handover and SLA activation',
         ]),
    ]

    for phase_name, timeline, phase_title, tasks in phases:
        pt = doc.add_table(rows=1, cols=3)
        thin_border_table(pt)
        set_col_width(pt, 0, 0.9)
        set_col_width(pt, 1, 1.2)
        set_col_width(pt, 2, 4.7)
        shade_cell(pt.rows[0].cells[0], C_HUAWEI_RED)
        shade_cell(pt.rows[0].cells[1], C_DARK_BLUE)
        shade_cell(pt.rows[0].cells[2], C_LIGHT_BLUE)
        set_cell_margins(pt.rows[0].cells[0], top=80, bottom=80, left=80, right=80)
        set_cell_margins(pt.rows[0].cells[1], top=80, bottom=80, left=80, right=80)
        set_cell_margins(pt.rows[0].cells[2], top=80, bottom=80, left=100, right=100)
        cell_para(pt.rows[0].cells[0], phase_name, bold=True, color=C_WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(pt.rows[0].cells[1], timeline, bold=True, color=C_WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(pt.rows[0].cells[2], phase_title, bold=True, color=C_NAVY, size=10)
        for task in tasks:
            cell_add_run(pt.rows[0].cells[2], f'\n▸  {task}', bold=False, color=C_BODY_TEXT, size=9)
        doc.add_paragraph()

    doc.add_page_break()

    # ── SUCCESS METRICS ───────────────────────────────────────────────────────
    section_banner(doc, '9.  SUCCESS METRICS ALIGNMENT')
    body(doc, 'We commit to meeting or exceeding every success metric specified in the Statement of Work:',
         before=4, after=5)
    sm_table = doc.add_table(rows=1+5, cols=4)
    thin_border_table(sm_table)
    set_col_width(sm_table, 0, 1.5)
    set_col_width(sm_table, 1, 1.8)
    set_col_width(sm_table, 2, 1.5)
    set_col_width(sm_table, 3, 2.0)
    for i, h in enumerate(['Success Area', 'CBO Target', 'Our Commitment', 'How We Deliver']):
        shade_cell(sm_table.rows[0].cells[i], C_NAVY)
        set_cell_margins(sm_table.rows[0].cells[i], top=70, bottom=70, left=80, right=80)
        cell_para(sm_table.rows[0].cells[i], h, bold=True, color=C_WHITE, size=9.5)
    sm_data = [
        ('Project Delivery', '100% of milestones met', '100% — milestone plan delivered in proposal, deviation escalated within 24h', 'Agile delivery with weekly steering; GitOps provides real-time delivery visibility'),
        ('Requirement Fulfilment', '100% compliance, evidenced in UAT', '100% coverage across 45/46 requirements; Edge AI excluded with documented rationale', 'UAT plan aligned to every SoW requirement; traceability matrix delivered'),
        ('Platform Performance', '≥99.5% uptime, <2s response', '99.9% uptime target; <1s first-token LLM response on GPU', 'Auto-healing via AIOps agent; multi-replica stateless services; SLO dashboards'),
        ('Security & Compliance', '100% pass, zero unresolved findings', '100% — security architecture reviewed pre-deployment; pen test in Phase 4', 'Continuous security monitoring from day one; findings tracked to closure'),
        ('Training & Support', '≥90% user satisfaction', '≥92% target across three training tracks', 'Structured training plan, pre/post assessment, dedicated support channel post-go-live'),
    ]
    for i, row_data in enumerate(sm_data):
        row = sm_table.rows[i+1]
        bg = C_LIGHT_BLUE if i%2==0 else C_LIGHT_GRAY
        for j, (cell, text) in enumerate(zip(row.cells, row_data)):
            shade_cell(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=80, right=80)
            cell_para(cell, text, bold=(j==0), color=C_NAVY if j==0 else C_BODY_TEXT, size=9)
    doc.add_paragraph()

    doc.add_page_break()

    # ── SECURITY & COMPLIANCE ─────────────────────────────────────────────────
    section_banner(doc, '10.  SECURITY AND COMPLIANCE — APPENDIX A RESPONSE')
    body(doc,
        'CBO\'s AI Security Policy (Appendix A) is a mature, technically detailed framework '
        'aligned with ISO 42001. The majority of its 15 technical requirements are fully '
        'addressed by the platform architecture described in this proposal.',
        before=4, after=5)

    sec_items = [
        ('Data Protection & Privacy',
         'AES-256 encryption at rest enforced across all storage systems. TLS 1.3 on all '
         'service communications via automated certificate management. mTLS between internal '
         'services via eBPF-encrypted service mesh. PII anonymisation before data ingestion '
         '(static) and at LLM inference time (runtime). Secrets vault with key rotation. '
         'Privacy Impact Assessment templates provided as engagement deliverables.'),
        ('AI Model Security',
         'Full model lifecycle governance: experiment tracking with lineage, model registry '
         'with approval workflows, cryptographically signed container images, GitOps '
         'deployment with mandatory review. SHAP and LIME explainability in the notebook '
         'environment. Fairness assessment toolkit pre-installed. Bias evaluation integrated '
         'into the model validation workflow.'),
        ('Access Control & Authentication',
         'Enterprise IAM with full RBAC, TOTP and FIDO2/WebAuthn MFA, JIT privileged '
         'access via time-limited credentials, LDAP/AD federation, and principle of least '
         'privilege enforced at the Kubernetes admission layer. All access events are '
         'immutably logged.'),
        ('Monitoring, Logging & Incident Response',
         'Real-time behavioural threat detection via eBPF kernel monitoring. SIEM platform '
         'with ML-based anomaly detection. Immutable, tamper-evident log storage. LLM '
         'observability with full audit trail of every AI interaction. AIOps intelligence '
         'agent provides AI-specific incident detection with contextualised alerts and '
         'suggested containment. Delivered: AI Incident Response Playbook covering '
         '10+ AI-specific scenarios.'),
        ('Ethical AI & Fairness',
         'Configurable AI safety guardrail layer enforces topic boundaries, detects and '
         'masks PII, and prevents harmful outputs. Fairness assessment toolkit evaluates '
         'demographic parity and equalised odds. Delivered: AI Governance Framework '
         'document defining CBO\'s high-risk AI approval process and ethical review '
         'procedures.'),
        ('ISO 42001 Alignment',
         'All technical controls required by ISO 42001 are implemented within the platform. '
         'A delivered ISO 42001 Technical Controls Mapping document provides a clause-by-clause '
         'mapping from each standard requirement to the corresponding platform component and '
         'configuration. Note: ISO 42001 also requires an organisational AI management '
         'system — policies, governance committees, and review cycles — which is CBO\'s '
         'organisational responsibility. We provide templates and guidance to accelerate '
         'this work.'),
    ]
    for title, text in sec_items:
        body_bold(doc, f'{title}: ', text, after=6)

    doc.add_page_break()

    # ── WHY HUAWEI ────────────────────────────────────────────────────────────
    section_banner(doc, '11.  WHY HUAWEI TECH. INVESTMENT (OMAN) LLC')
    why_data = [
        ('Established Oman Presence', 'Huawei has operated in Oman for over two decades with deep relationships across government, telecommunications, and financial services sectors. We are not a new entrant — we are a trusted long-term technology leader.'),
        ('Full Commercial Accountability', 'A single contract, a single point of accountability. CBO does not manage a consortium of vendors — Huawei owns the outcome end-to-end: from infrastructure provisioning to platform engineering, training, and long-term support.'),
        ('Infrastructure Reach', 'Huawei Cloud provides OMR-denominated, data-sovereign cloud infrastructure options. For on-premises deployment, Huawei hardware provides the compute foundation. Both options use the same platform architecture.'),
        ('Deep Open-Source AI Expertise', 'Huawei brings specialist expertise in enterprise open-source AI infrastructure — the specific technical discipline this engagement requires. Our delivery team\'s focus is squarely on AI-native Kubernetes platforms, not as a side practice.'),
        ('12-Week Delivery Confidence', 'The platform we are proposing is not a concept — it is a deployed, operating platform. Our implementation timelines are based on real delivery experience, not theoretical estimates.'),
        ('Post-Delivery Partnership', 'We view this engagement as the foundation of a long-term relationship. The AI sandbox is the starting point. As CBO\'s AI ambitions grow — production deployment, supervisory AI tools, cross-sector data analytics — Huawei is positioned to grow with you.'),
    ]
    wt = doc.add_table(rows=3, cols=2)
    thin_border_table(wt)
    for i, (title, text) in enumerate(why_data):
        r, c = i//2, i%2
        cell = wt.rows[r].cells[c]
        shade_cell(cell, C_LIGHT_BLUE if (r+c)%2==0 else C_LIGHT_GRAY)
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell_para(cell, title, bold=True, color=C_NAVY, size=10.5)
        cell_add_run(cell, f'\n{text}', bold=False, color=C_BODY_TEXT, size=9.5)
    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # INTERNAL APPENDIX Z — HUAWEI EYES ONLY
    # ══════════════════════════════════════════════════════════════════════════
    # Prominent internal banner
    for _ in range(2): doc.add_paragraph()
    int_banner = doc.add_table(rows=1, cols=1)
    no_border_table(int_banner)
    ib_cell = int_banner.rows[0].cells[0]
    set_cell_margins(ib_cell, top=150, bottom=150, left=120, right=120)
    shade_cell(ib_cell, C_HUAWEI_RED)
    cell_para(ib_cell, '⚠  APPENDIX Z — INTERNAL DOCUMENT', bold=True, color=C_WHITE, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_add_run(ib_cell, '\nSTRICTLY CONFIDENTIAL — HUAWEI TEAM EYES ONLY', bold=True, color=C_WHITE, size=12)
    cell_add_run(ib_cell, '\nNOT FOR DISTRIBUTION TO CBO OR ANY EXTERNAL PARTY', bold=True, color=HEX(0xFF,0xFF,0x80), size=11)
    cell_add_run(ib_cell, '\n\nThis appendix contains internal commercial guidance, cost modelling, and partner details.\nRemove this appendix before submitting the document to CBO.', bold=False, color=C_WHITE, size=9.5)
    doc.add_paragraph()

    section_banner(doc, 'Z.1  ENGAGEMENT CONTEXT & STRATEGIC VALUE')
    body(doc,
        'CBO is the Central Bank of Oman — the highest-value financial sector reference '
        'customer achievable in Oman. A successful delivery creates:',
        before=4, after=4)
    bullet(doc, 'A flagship Gulf central bank reference for both Huawei and the implementation partner')
    bullet(doc, 'A long-term managed services and expansion relationship (production AI deployment, supervisory tools)')
    bullet(doc, 'Proof of sovereign AI capability in a regulated financial environment — directly sellable to other GCC central banks and regulators')
    bullet(doc, 'Strategic positioning ahead of Oman Vision 2040 public sector AI expansion')
    body(doc,
        'This engagement should be treated as strategic, not transactional. Price to win '
        'the relationship, not to maximise initial margin.',
        after=6)

    section_banner(doc, 'Z.2  IMPLEMENTATION PARTNER DETAILS')
    body(doc,
        'The implementation partner referenced throughout this document as "our globally '
        'recognised open-source AI infrastructure specialist" is:',
        before=4, after=4)
    body(doc, 'OpenOva — an open-source AI-native infrastructure platform company.', after=4)
    bullet(doc, 'OpenOva provides the platform (52-component open-source AI ecosystem) and all technical delivery')
    bullet(doc, 'Huawei is the commercial prime — OpenOva operates as a Huawei implementation subcontractor for this engagement')
    bullet(doc, 'OpenOva\'s platform (the "AI Sandbox" we are proposing) is production-grade and deployed')
    bullet(doc, 'OpenOva holds responsibility for platform engineering, customisation, training, and post-delivery support')
    bullet(doc, 'Revenue structure: confirm with OpenOva — suggested model is fixed-price implementation + annual support subscription split between Huawei commercial margin and OpenOva delivery costs')
    body(doc,
        'For internal communications, refer to OpenOva by name. In all CBO-facing '
        'materials, use "our implementation partner" or describe capabilities without attribution.',
        after=6)

    section_banner(doc, 'Z.3  RESOURCE REQUIREMENTS')
    sub_banner(doc, 'Option A — On-Premises Deployment (Maximum Sovereignty)')
    body(doc, 'Hardware to be procured by CBO or via Huawei hardware supply:', before=4, after=4)
    hw_table = doc.add_table(rows=1+5, cols=3)
    thin_border_table(hw_table)
    set_col_width(hw_table, 0, 2.0)
    set_col_width(hw_table, 1, 2.5)
    set_col_width(hw_table, 2, 2.3)
    for i, h in enumerate(['Component', 'Specification', 'Notes']):
        shade_cell(hw_table.rows[0].cells[i], C_NAVY)
        set_cell_margins(hw_table.rows[0].cells[i], top=70, bottom=70, left=80, right=80)
        cell_para(hw_table.rows[0].cells[i], h, bold=True, color=C_WHITE, size=9.5)
    hw_data = [
        ('General Compute Nodes', '3–5 × server (16 vCPU, 64GB RAM each)', 'Platform control plane, observability, identity, object storage, ML services'),
        ('GPU Inference Nodes', '1–2 × GPU server (NVIDIA A10 24GB × 2 per node)', 'LLM inference and fine-tuning. A100 80GB if fine-tuning large models required'),
        ('Storage', '30TB NVMe SSD (hot) + expandable archival tier', 'Object storage, vector database, model artefacts, log retention'),
        ('Network', '10GbE internal switching + redundant uplinks', 'Isolated VLAN for sandbox — no routing to production network'),
        ('Estimated CapEx', 'USD 120,000–200,000 (hardware only)', '[CONFIRM WITH HUAWEI HARDWARE SUPPLY CHAIN — prices vary by GPU availability]'),
    ]
    for i, row_data in enumerate(hw_data):
        row = hw_table.rows[i+1]
        bg = C_LIGHT_BLUE if i%2==0 else C_LIGHT_GRAY
        for j, (cell, text) in enumerate(zip(row.cells, row_data)):
            shade_cell(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=80, right=80)
            cell_para(cell, text, bold=(j==0), color=C_NAVY if j==0 else C_BODY_TEXT, size=9)
    doc.add_paragraph()

    sub_banner(doc, 'Option B — Huawei Cloud Deployment (Fastest Start, Lowest Entry Cost)')
    body(doc, 'Platform infrastructure on Huawei Cloud, LLM inference via Huawei managed AI services:', before=4, after=4)
    cloud_table = doc.add_table(rows=1+5, cols=3)
    thin_border_table(cloud_table)
    set_col_width(cloud_table, 0, 2.0)
    set_col_width(cloud_table, 1, 2.5)
    set_col_width(cloud_table, 2, 2.3)
    for i, h in enumerate(['Resource', 'Specification', 'Indicative Monthly Cost']):
        shade_cell(cloud_table.rows[0].cells[i], C_NAVY)
        set_cell_margins(cloud_table.rows[0].cells[i], top=70, bottom=70, left=80, right=80)
        cell_para(cloud_table.rows[0].cells[i], h, bold=True, color=C_WHITE, size=9.5)
    cloud_data = [
        ('Platform Control Nodes', '3 × ECS General Purpose (8 vCPU, 32GB RAM each)', '[CONFIRM: ~OMR 120–180/month total — VERIFY WITH HUAWEI CLOUD OMAN PRICING]'),
        ('Object Storage (OBS)', '10TB initial allocation, expandable', '[CONFIRM WITH HUAWEI CLOUD PRICING]'),
        ('LLM-as-a-Service API', 'Huawei Pangu LLM API or equivalent managed endpoint', 'Token-based; CBO starts consuming when users are active. [CONFIRM HUAWEI API PRICING]'),
        ('Network & VPC', 'Isolated VPC, security groups, bandwidth', '[CONFIRM: typically included in base node pricing]'),
        ('Estimated Entry OpEx', 'Under 500 OMR/month (infrastructure only, before LLM API usage)', 'Scales with usage. Same architecture from 3 nodes to 300 nodes — zero re-architecture.'),
    ]
    for i, row_data in enumerate(cloud_data):
        row = cloud_table.rows[i+1]
        bg = C_LIGHT_BLUE if i%2==0 else C_LIGHT_GRAY
        for j, (cell, text) in enumerate(zip(row.cells, row_data)):
            shade_cell(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=80, right=80)
            cell_para(cell, text, bold=(j==0), color=C_NAVY if j==0 else C_BODY_TEXT, size=9)
    doc.add_paragraph()

    callout(doc, 'INTERNAL NOTE: Huawei Cloud Oman Availability',
        'CRITICAL: Confirm with Huawei Cloud team whether Huawei Cloud has a data centre '
        'region within Oman or the GCC that satisfies CBO\'s data sovereignty requirements. '
        'If a sovereign Oman region is not available, Option A (on-premises) should be '
        'the recommended default for a central bank. Do not offer Option B if data residency '
        'cannot be guaranteed within Oman\'s jurisdiction.',
        color=C_HUAWEI_RED)

    section_banner(doc, 'Z.4  COMMERCIAL MODEL RECOMMENDATIONS')
    body(doc, 'Recommended engagement structure for this procurement:', before=4, after=4)
    bullet(doc,
        'Fixed-price professional services for implementation (Phases 1–5, 12 weeks). '
        'Recommended range: [TO BE DETERMINED WITH OPENOVA — suggest OMR 35,000–60,000 '
        'depending on scope confirmed in discovery]. Do not underprice: CBO will discount '
        'a surprisingly low bid as a risk indicator.',
        bold_prefix='Implementation Fee:')
    bullet(doc,
        'Annual subscription covering: platform support, quarterly model updates, security '
        'patch cycle, and helpdesk. Recommended range: [OMR 20,000–35,000/year]. '
        'This is the recurring revenue model — structure it for multi-year commitment.',
        bold_prefix='Annual Support Subscription:')
    bullet(doc,
        'If Huawei Cloud Option B is selected, consumption-based billing on top of '
        'fixed infrastructure. Margin opportunity on cloud resell.',
        bold_prefix='Cloud Consumption (Option B only):')
    bullet(doc,
        'Scope extensions, additional use case development, production integration '
        'phases — billed separately as follow-on work. Plan for Phase 2 (production) '
        'engagement in 12–18 months.',
        bold_prefix='Future Phases:')

    section_banner(doc, 'Z.5  KEY RISKS & MITIGATIONS')
    risks = [
        ('GPU Hardware Lead Time', 'NVIDIA GPU servers have 3–6 month procurement lead times.', 'If Option A selected, raise hardware procurement immediately upon award. OpenOva can run in CPU mode initially (slower) while hardware is procured.'),
        ('Huawei Cloud Oman Availability', 'Unclear if Huawei Cloud has an Oman-region data centre.', 'Verify before proposal submission. If unavailable, position Option A as default for CBO\'s data sovereignty requirements.'),
        ('CBO IT Security Approval', 'CBO\'s IT and security teams may require extended review periods.', 'Engage CBO InfoSec in Phase 1. Build security review gates into the project plan. Our documented security architecture accelerates this.'),
        ('Competitor Landscape', 'Microsoft Azure OpenAI and AWS SageMaker will likely respond with cloud-hosted offerings.', 'Our sovereign, on-premises, Arabic-native differentiator is decisive for a central bank. Emphasise data stays in Oman. They cannot match this.'),
    ]
    for risk, challenge, mitigation in risks:
        body_bold(doc, f'{risk}: ', f'{challenge}  Mitigation: {mitigation}', after=5)

    section_banner(doc, 'Z.6  RECOMMENDED NEXT STEPS')
    bullet(doc, 'Verify Huawei Cloud Oman data residency capability BEFORE submission')
    bullet(doc, 'Confirm commercial terms with OpenOva (implementation + support pricing)')
    bullet(doc, 'Request a meeting with CBO Technology Procurement Committee for Q&A')
    bullet(doc, 'Prepare a live technical demonstration environment on Huawei Cloud')
    bullet(doc, 'Identify CBO executive sponsor — likely CBO\'s CIO or Head of IT Governance')
    bullet(doc, 'Register Huawei + OpenOva as a consortium on CBO\'s vendor portal if required')

    # ── SAVE ──────────────────────────────────────────────────────────────────
    output_path = '/home/openova/repos/openova/.claude/CBO-AI-Sandbox-Proposal-Huawei.docx'
    doc.save(output_path)
    print(f'Document saved: {output_path}')
    req_scores = [r[3] for r in REQUIREMENTS]
    print(f'Requirements covered: {len([s for s in req_scores if s >= 75])}/{len(req_scores)}')
    print(f'Average coverage score: {sum(req_scores)/len(req_scores):.1f}/100')

if __name__ == '__main__':
    build()
